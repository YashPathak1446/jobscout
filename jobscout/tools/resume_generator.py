"""
Resume Generator — Tailors resume per JD and outputs .docx

Two-step process:
1. LLM (Gemini) rewrites resume bullets using XYZ formula, mirrors JD language
2. docx builder formats the tailored content into a clean .docx file

The LLM receives: master resume, JD text, selected components, and config rules.
It returns structured JSON that the docx builder consumes.
"""

import os
import json
import logging
import subprocess
import tempfile
from pathlib import Path

logger = logging.getLogger(__name__)


# =========================================================================
# LLM RESUME TAILORING
# =========================================================================

def tailor_resume_with_llm(
    master_resume_text: str,
    jd_text: str,
    selected_experience_ids: list[str],
    selected_project_ids: list[str],
    lead_skills: list[str],
    parsed_resume,
    resume_rules: str,
    similar_tech_map: dict,
    model: str = "gemini-2.5-flash",
    fallback_model: str = "gemini-2.5-flash",
) -> dict | None:
    """
    Use Gemini to tailor resume content for a specific JD.

    Returns structured dict with tailored sections, or None on failure.
    """
    # Build component context
    selected_exp_text = ""
    for exp in parsed_resume.experiences:
        if exp.id in selected_experience_ids:
            selected_exp_text += f"\n--- {exp.title} @ {exp.organization} ({exp.date_range}) ---\n"
            selected_exp_text += f"Tech: {exp.tech_line}\n"
            for bullet in exp.bullets:
                selected_exp_text += f"- {bullet}\n"

    selected_proj_text = ""
    for proj in parsed_resume.projects:
        if proj.id in selected_project_ids:
            selected_proj_text += f"\n--- {proj.title} ({proj.date_range}) ---\n"
            selected_proj_text += f"Tech: {proj.tech_line}\n"
            for bullet in proj.bullets:
                selected_proj_text += f"- {bullet}\n"

    prompt = f"""You are an expert resume writer specializing in ATS-optimized resumes.
Your task is to tailor resume content for a specific job description, following the EXACT format below.

CRITICAL RULES — FOLLOW EXACTLY:
{resume_rules}

BULLET COUNT RULES (STRICT):
- Work experiences: 3-4 bullets MAXIMUM per role
- Projects: 2-3 bullets MAXIMUM per project
- Select only the most relevant and impactful bullets
- Never repeat or pad bullets

BULLET QUALITY RULES:
- Use XYZ formula: "Accomplished [X], as measured by [Y], by doing [Z]"
- Strong action verbs: Architected, Engineered, Optimized, Deployed, Implemented
- Mirror exact JD terminology
- Keep all metrics exactly as in original — never invent numbers
- Never fabricate skills, tools, or experiences

SIMILAR TECHNOLOGY NOTE:
If JD mentions a technology the candidate doesn't have but has a similar one,
include the project with the ACTUAL technology. Never claim skills they don't have.

=== JOB DESCRIPTION ===
{jd_text[:5000]}

=== SELECTED EXPERIENCES (pick 3-4 bullets each, most relevant to JD) ===
{selected_exp_text}

=== SELECTED PROJECTS (pick 2-3 bullets each, most relevant to JD) ===
{selected_proj_text}

=== CANDIDATE SKILLS (reorder to lead with JD-matched skills) ===
{parsed_resume.skills_text}

=== JD-MATCHED LEAD SKILLS ===
{', '.join(lead_skills[:10]) if lead_skills else 'Use your judgment based on JD'}

Output ONLY valid JSON (no markdown, no backticks):
{{
    "skills": {{
        "Languages": "Python, Java, ...",
        "Cloud & Infrastructure": "AWS, Docker, ...",
        "Databases & Search": "MongoDB, ...",
        "Frameworks & Libraries": "Flask, Node.js, ...",
        "AI & Data Science": "PyTorch, ...",
        "Developer Tools": "Git, CI/CD, ..."
    }},
    "experiences": [
        {{
            "title": "Software Engineer Intern",
            "company": "Sorenson Communications",
            "location": "Salt Lake City, UT",
            "dates": "June 2025 -- Oct. 2025",
            "bullets": [
                "Bullet 1 (XYZ format)",
                "Bullet 2 (XYZ format)",
                "Bullet 3 (XYZ format)"
            ]
        }}
    ],
    "projects": [
        {{
            "name": "JobScout - AI Job Automation",
            "url": "https://github.com/YashPathak1446/jobscout",
            "tech": "Python, Google ADK, Gemini API, Docker",
            "dates": "Jan. 2026 - Current",
            "bullets": [
                "Bullet 1 (XYZ format)",
                "Bullet 2 (XYZ format)"
            ]
        }}
    ]
}}
"""

    # Try primary model, fallback on error
    for attempt_model in [model, fallback_model]:
        try:
            from google import genai

            client = genai.Client(api_key=os.getenv("GOOGLE_API_KEY"))
            response = client.models.generate_content(
                model=attempt_model,
                contents=prompt,
            )

            raw_text = response.text.strip()
            # Strip markdown code fences if present
            if raw_text.startswith("```"):
                raw_text = raw_text.split("\n", 1)[1] if "\n" in raw_text else raw_text[3:]
            if raw_text.endswith("```"):
                raw_text = raw_text[:-3]
            raw_text = raw_text.strip()
            if raw_text.startswith("json"):
                raw_text = raw_text[4:].strip()

            data = json.loads(raw_text)
            logger.info(f"Resume tailored with {attempt_model}")
            return data

        except json.JSONDecodeError as e:
            logger.error(f"JSON parse error from {attempt_model}: {e}")
            logger.debug(f"Raw response: {raw_text[:500]}")
            if attempt_model == fallback_model:
                return None
        except Exception as e:
            error_str = str(e).lower()
            if "429" in error_str or "rate limit" in error_str:
                logger.warning(f"{attempt_model} rate limited, trying fallback")
                continue
            logger.error(f"LLM error with {attempt_model}: {e}")
            if attempt_model == fallback_model:
                return None

    return None


def tailor_resume_mock(
    parsed_resume,
    selected_experience_ids: list[str],
    selected_project_ids: list[str],
    lead_skills: list[str],
) -> dict:
    """Mock tailoring for testing -- returns original content without LLM."""
    # Build a quick location lookup from raw LaTeX if available
    location_map = {}
    raw_tex = getattr(parsed_resume, 'raw_text', '')
    if raw_tex:
        import re as _re
        for m in _re.finditer(
            r'\\resumeSubheading\s*\{[^}]*\}\{[^}]*\}\s*\{([^}]*)\}\{([^}]*)\}',
            raw_tex
        ):
            company = m.group(1).strip()
            location = m.group(2).strip()
            if company and location:
                location_map[company.lower()] = location

    experiences = []
    for exp in parsed_resume.experiences:
        if exp.id in selected_experience_ids:
            title_str = exp.title
            org = exp.organization or ""
            location = location_map.get(org.lower(), "")
            experiences.append({
                "title": title_str,
                "company": org,
                "location": location,
                "dates": exp.date_range,
                "bullets": exp.bullets[:4],
            })

    projects = []
    for proj in parsed_resume.projects:
        if proj.id in selected_project_ids:
            # Extract URL from raw LaTeX source
            url = ""
            import re as _re
            raw_tex = getattr(parsed_resume, 'raw_text', '')
            if raw_tex and proj.title:
                name_chunk = proj.title[:12]
                import re as _re
                for line in raw_tex.split('\n'):
                    if name_chunk in line and 'href' in line:
                        m = _re.search(r'href\{([^}]+)\}', line)
                        if m:
                            url = m.group(1)
                            break
            projects.append({
                "name": proj.title,
                "url": url,
                "tech": proj.tech_line,
                "dates": proj.date_range,
                "bullets": proj.bullets[:3],
            })

    # Parse contact info from resume
    contact_lines = parsed_resume.contact_info.split("\n")
    name = contact_lines[0].strip() if contact_lines else "Yash Pathak"

    return {
        "contact": {
            "name": "Yash Pathak",
            "location": "Tracy, CA",
            "email": "REDACTED-EMAIL",
            "phone": "REDACTED-PHONE",
            "links": [
                "https://github.com/YashPathak1446",
                "https://www.linkedin.com/in/yash-pathak-6074a1210/",
            ],
            "summary": "",
        },
        "education": {
            "school": "University of California, Irvine",
            "location": "Irvine, CA",
            "degree": "Bachelor of Science in Computer Science",
            "dates": "Sep. 2021 -- June 2025",
            "details": "Analysis of Algorithms, Database Systems, Web Applications, AI/ML, System Design, Computer Networks",
        },
        "skills": {
            "Languages": "Python, Java, C++, SQL (PostgreSQL, MySQL), JavaScript/TypeScript, HTML/CSS",
            "Cloud & Infrastructure": "AWS (EC2, S3, Lambda, API Gateway), Docker, Kubernetes, Terraform, Nginx, Firebase",
            "Databases & Search": "MongoDB, Vector Databases (Weaviate), Inverted Indexing, RAG",
            "Frameworks & Libraries": "Flask, Node.js, Express, Angular, BeautifulSoup, lxml",
            "AI & Data Science": "PyTorch, TensorFlow, Scikit-learn, Pandas, NumPy, Matplotlib",
            "Developer Tools": "Git, GitHub Actions, CI/CD, Postman, Linux/Unix Shell, Figma",
        },
        "experiences": experiences,
        "projects": projects,
    }


# =========================================================================
# DOCX BUILDER
# =========================================================================

def build_resume_docx(tailored_data: dict, output_path: str) -> str:
    """
    Build a .docx resume from tailored data using docx-js (Node.js).
    Returns the path to the generated file.
    """
    # Generate the JS script
    js_code = _generate_docx_js(tailored_data, output_path)

    # Write to temp file and run
    js_path = os.path.join(tempfile.gettempdir(), "build_resume.mjs")
    with open(js_path, "w", encoding="utf-8") as f:
        f.write(js_code)

    try:
        result = subprocess.run(
            ["node", js_path],
            capture_output=True, text=True, timeout=30,
        )
        if result.returncode != 0:
            logger.error(f"docx-js error: {result.stderr}")
            # Fallback to python-docx
            return _build_resume_python_docx(tailored_data, output_path)
        logger.info(f"Resume saved: {output_path}")
        return output_path
    except FileNotFoundError:
        logger.warning("Node.js not found, falling back to python-docx")
        return _build_resume_python_docx(tailored_data, output_path)
    except Exception as e:
        logger.error(f"docx build error: {e}")
        return _build_resume_python_docx(tailored_data, output_path)


def _generate_docx_js(data: dict, output_path: str) -> str:
    """Generate the JavaScript code for docx-js resume builder."""
    contact = data.get("contact", {})
    edu = data.get("education", {})
    skills = data.get("skills", {})
    experiences = data.get("experiences", [])
    projects = data.get("projects", [])

    # Escape strings for JS
    def esc(s):
        return str(s).replace("\\", "\\\\").replace('"', '\\"').replace("\n", " ")

    # Build experience sections
    exp_children = ""
    for exp in experiences:
        bullets_js = ""
        for bullet in exp.get("bullets", []):
            bullets_js += f"""
        new Paragraph({{
            numbering: {{ reference: "bullets", level: 0 }},
            spacing: {{ after: 40 }},
            children: [new TextRun({{ text: "{esc(bullet)}", size: 20, font: "Arial" }})]
        }}),"""

        exp_children += f"""
        new Paragraph({{
            spacing: {{ before: 160, after: 40 }},
            children: [
                new TextRun({{ text: "{esc(exp.get('title', ''))}", bold: true, size: 21, font: "Arial" }}),
                new TextRun({{ text: " | {esc(exp.get('company', ''))}", size: 21, font: "Arial" }}),
            ]
        }}),
        new Paragraph({{
            spacing: {{ after: 60 }},
            children: [
                new TextRun({{ text: "{esc(exp.get('dates', ''))}", italics: true, size: 20, font: "Arial", color: "666666" }}),
            ]
        }}),{bullets_js}"""

    # Build project sections
    proj_children = ""
    for proj in projects:
        bullets_js = ""
        for bullet in proj.get("bullets", []):
            bullets_js += f"""
        new Paragraph({{
            numbering: {{ reference: "bullets", level: 0 }},
            spacing: {{ after: 40 }},
            children: [new TextRun({{ text: "{esc(bullet)}", size: 20, font: "Arial" }})]
        }}),"""

        tech_str = f" | {esc(proj.get('tech', ''))}" if proj.get('tech') else ""
        proj_children += f"""
        new Paragraph({{
            spacing: {{ before: 160, after: 40 }},
            children: [
                new TextRun({{ text: "{esc(proj.get('name', ''))}", bold: true, size: 21, font: "Arial" }}),
                new TextRun({{ text: "{tech_str}", size: 20, font: "Arial" }}),
            ]
        }}),
        new Paragraph({{
            spacing: {{ after: 60 }},
            children: [
                new TextRun({{ text: "{esc(proj.get('dates', ''))}", italics: true, size: 20, font: "Arial", color: "666666" }}),
            ]
        }}),{bullets_js}"""

    # Build skills lines
    skills_children = ""
    for category, value in skills.items():
        if value and value.strip():
            label = category.replace("_", " ").title()
            skills_children += f"""
        new Paragraph({{
            spacing: {{ after: 40 }},
            children: [
                new TextRun({{ text: "{esc(label)}: ", bold: true, size: 20, font: "Arial" }}),
                new TextRun({{ text: "{esc(value)}", size: 20, font: "Arial" }}),
            ]
        }}),"""

    # Links
    links_text = " | ".join(contact.get("links", []))

    abs_output = os.path.abspath(output_path).replace("\\", "/")

    return f"""
const {{ Document, Packer, Paragraph, TextRun, AlignmentType,
         HeadingLevel, LevelFormat, BorderStyle }} = require("docx");
const fs = require("fs");

const doc = new Document({{
    numbering: {{
        config: [{{
            reference: "bullets",
            levels: [{{
                level: 0,
                format: LevelFormat.BULLET,
                text: "\\u2022",
                alignment: AlignmentType.LEFT,
                style: {{ paragraph: {{ indent: {{ left: 360, hanging: 180 }} }} }}
            }}]
        }}]
    }},
    styles: {{
        default: {{ document: {{ run: {{ font: "Arial", size: 22 }} }} }},
    }},
    sections: [{{
        properties: {{
            page: {{
                size: {{ width: 12240, height: 15840 }},
                margin: {{ top: 720, right: 720, bottom: 720, left: 720 }}
            }}
        }},
        children: [
            // Name
            new Paragraph({{
                alignment: AlignmentType.CENTER,
                spacing: {{ after: 40 }},
                children: [new TextRun({{ text: "{esc(contact.get('name', ''))}", bold: true, size: 28, font: "Arial" }})]
            }}),
            // Contact line
            new Paragraph({{
                alignment: AlignmentType.CENTER,
                spacing: {{ after: 40 }},
                children: [new TextRun({{ text: "{esc(contact.get('email', ''))} | {esc(contact.get('phone', ''))}", size: 20, font: "Arial" }})]
            }}),
            // Links
            new Paragraph({{
                alignment: AlignmentType.CENTER,
                spacing: {{ after: 120 }},
                border: {{ bottom: {{ style: BorderStyle.SINGLE, size: 6, color: "999999" }} }},
                children: [new TextRun({{ text: "{esc(links_text)}", size: 20, font: "Arial", color: "0563C1" }})]
            }}),

            // EDUCATION
            new Paragraph({{
                spacing: {{ before: 120, after: 60 }},
                border: {{ bottom: {{ style: BorderStyle.SINGLE, size: 4, color: "CCCCCC" }} }},
                children: [new TextRun({{ text: "EDUCATION", bold: true, size: 22, font: "Arial" }})]
            }}),
            new Paragraph({{
                spacing: {{ after: 40 }},
                children: [
                    new TextRun({{ text: "{esc(edu.get('school', ''))}", bold: true, size: 21, font: "Arial" }}),
                    new TextRun({{ text: " | {esc(edu.get('degree', ''))}", size: 21, font: "Arial" }}),
                    new TextRun({{ text: " | {esc(edu.get('dates', ''))}", size: 20, font: "Arial", color: "666666" }}),
                ]
            }}),
            new Paragraph({{
                spacing: {{ after: 60 }},
                children: [new TextRun({{ text: "{esc(edu.get('details', ''))}", size: 20, font: "Arial" }})]
            }}),

            // SKILLS
            new Paragraph({{
                spacing: {{ before: 120, after: 60 }},
                border: {{ bottom: {{ style: BorderStyle.SINGLE, size: 4, color: "CCCCCC" }} }},
                children: [new TextRun({{ text: "TECHNICAL SKILLS", bold: true, size: 22, font: "Arial" }})]
            }}),
            {skills_children}

            // EXPERIENCE
            new Paragraph({{
                spacing: {{ before: 120, after: 60 }},
                border: {{ bottom: {{ style: BorderStyle.SINGLE, size: 4, color: "CCCCCC" }} }},
                children: [new TextRun({{ text: "EXPERIENCE", bold: true, size: 22, font: "Arial" }})]
            }}),
            {exp_children}

            // PROJECTS
            new Paragraph({{
                spacing: {{ before: 120, after: 60 }},
                border: {{ bottom: {{ style: BorderStyle.SINGLE, size: 4, color: "CCCCCC" }} }},
                children: [new TextRun({{ text: "PROJECTS", bold: true, size: 22, font: "Arial" }})]
            }}),
            {proj_children}
        ]
    }}]
}});

Packer.toBuffer(doc).then(buffer => {{
    fs.writeFileSync("{abs_output}", buffer);
    console.log("Resume saved: {abs_output}");
}});
"""


def _build_resume_python_docx(data: dict, output_path: str) -> str:
    """Build resume with python-docx matching the candidate actual resume format."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        logger.error("python-docx not installed. pip install python-docx")
        return ""

    doc = Document()

    # Page setup — tight margins like the example resume
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    def add_bottom_border(paragraph):
        """Add a bottom border line to a paragraph (section divider)."""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def set_spacing(paragraph, before=0, after=0):
        """Set paragraph spacing in points."""
        paragraph.paragraph_format.space_before = Pt(before)
        paragraph.paragraph_format.space_after = Pt(after)

    contact = data.get("contact", {})
    edu = data.get("education", {})
    skills = data.get("skills", {})

    # ===== NAME (bold, centered, large) =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=0)
    run = p.add_run(contact.get("name", "YASH BALAJI PATHAK").upper())
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = "Arial"

    # ===== CONTACT LINE (centered, with embedded links) =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=2)

    # Build contact: location | email | GitHub | LinkedIn | phone
    contact_parts = []
    location = contact.get("location", "Tracy, CA (Bay Area)")
    if location:
        contact_parts.append(location)
    email = contact.get("email", "")
    if email and email != "email@example.com":
        contact_parts.append(email)
    phone = contact.get("phone", "")
    if phone and phone != "555-555-5555":
        contact_parts.append(phone)

    # Add links as text (python-docx hyperlinks are complex)
    links = contact.get("links", [])
    for link in links:
        if "github" in link.lower():
            contact_parts.append("GitHub")
        elif "linkedin" in link.lower():
            contact_parts.append("LinkedIn")
        elif "portfolio" in link.lower():
            contact_parts.append("portfolio")
        else:
            contact_parts.append(link)

    run = p.add_run(" | ".join(contact_parts))
    run.font.size = Pt(9)
    run.font.name = "Arial"

    # ===== PROFESSIONAL SUMMARY (if provided) =====
    summary = contact.get("summary", "")
    if summary:
        p = doc.add_paragraph()
        set_spacing(p, before=4, after=4)
        run = p.add_run(summary)
        run.font.size = Pt(9)
        run.font.name = "Arial"

    # ===== EDUCATION =====
    p = doc.add_paragraph()
    set_spacing(p, before=6, after=2)
    add_bottom_border(p)
    run = p.add_run("EDUCATION")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Arial"

    # School line: bold school name, then dates right-aligned via tab
    p = doc.add_paragraph()
    set_spacing(p, before=2, after=0)
    run = p.add_run(edu.get("school", "University of California - Irvine"))
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Arial"
    run = p.add_run(f"\t\t\t\t\t{edu.get('dates', '')}")
    run.font.size = Pt(10)
    run.font.name = "Arial"
    run.italic = True

    # Degree
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=0)
    run = p.add_run(edu.get("degree", "Bachelor of Science in Computer Science"))
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Arial"

    # Courses (if provided)
    details = edu.get("details", "")
    if details:
        p = doc.add_paragraph()
        set_spacing(p, before=0, after=0)
        run = p.add_run("Courses: ")
        run.bold = True
        run.italic = True
        run.font.size = Pt(9)
        run.font.name = "Arial"
        run = p.add_run(details)
        run.italic = True
        run.font.size = Pt(9)
        run.font.name = "Arial"

    # ===== TECHNICAL SKILLS =====
    p = doc.add_paragraph()
    set_spacing(p, before=6, after=2)
    add_bottom_border(p)
    run = p.add_run("TECHNICAL SKILLS")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Arial"

    for category, value in skills.items():
        if value and value.strip():
            label = category.replace("_", " & ").replace("and", "&").title()
            # Match the resume format: "Programming & Core:"
            p = doc.add_paragraph()
            set_spacing(p, before=0, after=0)
            run = p.add_run(f"{label}: ")
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = "Arial"
            run = p.add_run(value)
            run.font.size = Pt(10)
            run.font.name = "Arial"

    # ===== EXPERIENCE =====
    p = doc.add_paragraph()
    set_spacing(p, before=6, after=2)
    add_bottom_border(p)
    run = p.add_run("EXPERIENCE")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Arial"

    for exp in data.get("experiences", []):
        # Title line: "Software Engineer Intern, Company" + right-aligned dates
        p = doc.add_paragraph()
        set_spacing(p, before=4, after=0)
        title_str = exp.get("title", "")
        company_str = exp.get("company", "")
        if company_str and company_str not in title_str:
            title_str = f"{title_str}, {company_str}"
        run = p.add_run(title_str)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Arial"
        dates = exp.get("dates", "")
        if dates:
            run = p.add_run(f"\t\t\t\t{dates}")
            run.font.size = Pt(10)
            run.font.name = "Arial"
            run.italic = True

        # Subtitle (if present)
        subtitle = exp.get("subtitle", "")
        if subtitle:
            p = doc.add_paragraph()
            set_spacing(p, before=0, after=0)
            run = p.add_run(subtitle)
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = "Arial"

        # Bullets
        for bullet in exp.get("bullets", []):
            p = doc.add_paragraph(style="List Bullet")
            set_spacing(p, before=0, after=1)
            # Clear default and add custom run
            p.clear()
            run = p.add_run(bullet)
            run.font.size = Pt(9.5)
            run.font.name = "Arial"

        # Tech line
        tech = exp.get("tech", "")
        if tech:
            p = doc.add_paragraph()
            set_spacing(p, before=1, after=0)
            run = p.add_run("Tech: ")
            run.bold = True
            run.font.size = Pt(9.5)
            run.font.name = "Arial"
            run = p.add_run(tech)
            run.font.size = Pt(9.5)
            run.font.name = "Arial"

    # ===== PROJECTS =====
    p = doc.add_paragraph()
    set_spacing(p, before=6, after=2)
    add_bottom_border(p)
    run = p.add_run("ACADEMIC PROJECTS")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Arial"

    for proj in data.get("projects", []):
        # Project header: "Name | Tech" + right-aligned dates
        p = doc.add_paragraph()
        set_spacing(p, before=4, after=0)
        run = p.add_run(proj.get("name", ""))
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Arial"
        tech = proj.get("tech", "")
        if tech:
            run = p.add_run(f" | ")
            run.font.size = Pt(10)
            run.font.name = "Arial"
            run = p.add_run(tech)
            run.italic = True
            run.font.size = Pt(9.5)
            run.font.name = "Arial"
        dates = proj.get("dates", "")
        if dates:
            run = p.add_run(f"\t\t{dates}")
            run.font.size = Pt(9.5)
            run.font.name = "Arial"
            run.italic = True

        # Bullets
        for bullet in proj.get("bullets", []):
            p = doc.add_paragraph(style="List Bullet")
            set_spacing(p, before=0, after=1)
            p.clear()
            run = p.add_run(bullet)
            run.font.size = Pt(9.5)
            run.font.name = "Arial"

    # Save
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    logger.info(f"Resume saved (python-docx): {output_path}")
    return output_path


# =========================================================================
# LATEX BUILDER — uses main.tex as template
# =========================================================================

def _escape_latex(text: str) -> str:
    """Escape special LaTeX characters."""
    chars = {
        '&': r'\&', '%': r'\%', '$': r'\$', '#': r'\#',
        '_': r'\_', '{': r'\{', '}': r'\}', '~': r'\textasciitilde{}',
        '^': r'\textasciicircum{}', '<': r'\textless{}', '>': r'\textgreater{}',
    }
    for char, replacement in chars.items():
        text = text.replace(char, replacement)
    return text


def _build_experience_latex(experiences: list[dict]) -> str:
    """Build LaTeX for experience section using resumeSubheading format."""
    esc = _escape_latex
    latex = ""
    for exp in experiences:
        title = esc(exp.get("title", ""))
        company = esc(exp.get("company", ""))
        location = esc(exp.get("location", ""))
        dates = esc(exp.get("dates", ""))

        latex += f"""    \\resumeSubheading
      {{{title}}}{{{dates}}}
      {{{company}}}{{{location}}}
      \\resumeItemListStart\n"""

        for bullet in exp.get("bullets", []):
            latex += f"        \\resumeItem{{{esc(bullet)}}}\n"

        latex += "      \\resumeItemListEnd\n\n"
    return latex


def _build_project_latex(projects: list[dict]) -> str:
    """Build LaTeX for projects section using resumeProjectHeading format."""
    esc = _escape_latex
    latex = ""
    for proj in projects:
        name = proj.get("name", "")
        url = proj.get("url", "")
        tech = esc(proj.get("tech", ""))
        dates = esc(proj.get("dates", ""))

        # Create hyperlinked name if URL provided
        if url:
            name_latex = f"\\textbf{{\\href{{{url}}}{{\\underline{{{esc(name)}}}}}}}"
        else:
            name_latex = f"\\textbf{{{esc(name)}}}"

        if tech:
            heading = f"{name_latex} $|$ \\emph{{{tech}}}"
        else:
            heading = name_latex

        latex += f"""      \\resumeProjectHeading
          {{{heading}}}{{{dates}}}
          \\resumeItemListStart\n"""

        for bullet in proj.get("bullets", []):
            latex += f"            \\resumeItem{{{esc(bullet)}}}\n"

        latex += "          \\resumeItemListEnd\n\n"
    return latex


def _build_skills_latex(skills: dict) -> str:
    """Build LaTeX for technical skills section."""
    esc = _escape_latex
    lines = ""
    for label, value in skills.items():
        if value and value.strip():
            lines += f"     \\textbf{{{esc(label)}}}{{: {esc(value)}}} \\\\\n"
    return lines


def build_resume_latex(tailored_data: dict, output_path: str, template_path: str = None) -> str:
    """
    Build a resume PDF using the candidate's own main.tex as a template.
    Only replaces Experience, Projects, and Technical Skills sections.
    Everything else (header, education, formatting) stays exactly as-is.

    Returns path to the generated PDF (or .tex if pdflatex unavailable).
    """
    # Find the template
    if template_path is None:
        candidates = [
            "data/master_resume.tex",
            os.path.join(os.path.dirname(__file__), "../../data/master_resume.tex"),
        ]
        for c in candidates:
            if os.path.exists(c):
                template_path = c
                break

    if template_path and os.path.exists(template_path):
        with open(template_path, "r", encoding="utf-8") as f:
            template = f.read()
        logger.info(f"Using template: {template_path}")
    else:
        # Fall back to generating from scratch if no template
        logger.warning("No template found, generating from scratch")
        template = None

    experiences = tailored_data.get("experiences", [])
    projects = tailored_data.get("projects", [])
    skills = tailored_data.get("skills", {})

    exp_latex = _build_experience_latex(experiences)
    proj_latex = _build_project_latex(projects)
    skills_latex = _build_skills_latex(skills)

    if template:
        import re

        # Match \resumeSubHeadingListEnd only on lines not starting with %
        # We do this by pre-processing: temporarily mark the real (non-commented) end
        # with a unique token, then use that in the regex
        MARKER = "XRESUMESUBHEADINGLISTENDX"
        marked = re.sub(
            r'^(\s*)\\resumeSubHeadingListEnd',
            rf'\1{MARKER}',
            template,
            flags=re.MULTILINE,
        )
        # Remove markers that were on commented lines (they won't have % prefix,
        # but commented \resumeSubHeadingListEnd lines start with %)
        # Actually we already handled this: only non-commented lines get the marker
        # because the regex ^(\s*)\\resumeSubHeadingListEnd matches lines where
        # the only thing before \resumeSubHeadingListEnd is whitespace

        def replace_exp(m):
            return m.group(1) + "\n\n" + exp_latex + "  " + m.group(3)

        def replace_proj(m):
            return m.group(1) + "\n" + proj_latex + "    " + m.group(3)

        def replace_skills(m):
            return m.group(1) + "\n" + skills_latex + "    " + m.group(3)

        # Replace Experience section
        marked = re.sub(
            r'(\\section\{Experience\}\s*\\resumeSubHeadingListStart)(.*?)(\s*' + MARKER + r')',
            replace_exp, marked, flags=re.DOTALL, count=1,
        )

        # Replace Projects section
        marked = re.sub(
            r'(\\section\{Projects\}\s*\\resumeSubHeadingListStart)(.*?)(\s*' + MARKER + r')',
            replace_proj, marked, flags=re.DOTALL, count=1,
        )

        # Restore remaining markers to original command
        latex = marked.replace(MARKER, r'\resumeSubHeadingListEnd')

        # Replace Technical Skills — use string split instead of regex to avoid escape issues
        skills_start = r'\section{Technical Skills}'
        item_start = r'\small{\item{'
        item_end = r'}}'
        if skills_start in latex and item_start in latex:
            sec_idx = latex.find(skills_start)
            item_idx = latex.find(item_start, sec_idx)
            if item_idx > 0:
                item_content_start = item_idx + len(item_start)
                # Find the closing }}
                close_idx = latex.find(item_end, item_content_start)
                if close_idx > 0:
                    latex = (
                        latex[:item_content_start] +
                        "\n" + skills_latex + "    " +
                        latex[close_idx:]
                    )
    else:
        # Generate full LaTeX from scratch (Jake's template)
        latex = _generate_full_latex(tailored_data, exp_latex, proj_latex, skills_latex)

    # Write .tex and compile
    tex_dir = os.path.dirname(output_path) or "."
    os.makedirs(tex_dir, exist_ok=True)
    base = os.path.splitext(output_path)[0]
    tex_path = base + ".tex"
    pdf_path = base + ".pdf"

    with open(tex_path, "w", encoding="utf-8") as f:
        f.write(latex)

    try:
        result = subprocess.run(
            ["pdflatex", "-interaction=nonstopmode", "-output-directory", tex_dir, tex_path],
            capture_output=True, text=True, timeout=30,
        )
        if os.path.exists(pdf_path):
            for ext in [".aux", ".log", ".out"]:
                aux = base + ext
                if os.path.exists(aux):
                    try:
                        os.remove(aux)
                    except Exception:
                        pass
            logger.info(f"Resume PDF saved: {pdf_path}")
            return pdf_path
        else:
            logger.error(f"pdflatex error: {result.stderr[-300:]}")
            print(f"\n    ⚠️  pdflatex could not compile. .tex file saved at:\n    {tex_path}")
            print(f"    → Compile on Overleaf: overleaf.com → New Project → paste .tex content")
            return tex_path
    except FileNotFoundError:
        logger.warning("pdflatex not found. Saving .tex file only.")
        return tex_path
    except Exception as e:
        logger.error(f"LaTeX compile error: {e}")
        return tex_path


def _generate_full_latex(data: dict, exp_latex: str, proj_latex: str, skills_latex: str) -> str:
    """Generate a complete LaTeX document when no template is available."""
    contact = data.get("contact", {})
    edu = data.get("education", {})
    esc = _escape_latex

    name = esc(contact.get("name", "Yash Pathak"))
    email = contact.get("email", "")
    phone = contact.get("phone", "")
    links = contact.get("links", [])

    contact_parts = []
    if phone:
        contact_parts.append(f"\\small {esc(phone)}")
    if email:
        contact_parts.append(f"\\href{{mailto:{email}}}{{\\underline{{{esc(email)}}}}}")
    for link in links:
        if "github" in link.lower():
            contact_parts.append(f"\\href{{{link}}}{{\\underline{{GitHub}}}}")
        elif "linkedin" in link.lower():
            contact_parts.append(f"\\href{{{link}}}{{\\underline{{LinkedIn}}}}")

    contact_line = " $|$ ".join(contact_parts)

    edu_school = esc(edu.get("school", "University of California, Irvine"))
    edu_location = esc(edu.get("location", "Irvine, CA"))
    edu_degree = esc(edu.get("degree", "Bachelor of Science in Computer Science"))
    edu_dates = esc(edu.get("dates", "Sep. 2021 -- June 2025"))
    edu_courses = esc(edu.get("details", ""))

    return f"""\\documentclass[letterpaper,11pt]{{article}}
\\usepackage{{latexsym}}
\\usepackage[empty]{{fullpage}}
\\usepackage{{titlesec}}
\\usepackage{{marvosym}}
\\usepackage[usenames,dvipsnames]{{color}}
\\usepackage{{verbatim}}
\\usepackage{{enumitem}}
\\usepackage[hidelinks]{{hyperref}}
\\usepackage{{fancyhdr}}
\\usepackage[english]{{babel}}
\\usepackage{{tabularx}}
\\input{{glyphtounicode}}
\\pagestyle{{fancy}}
\\fancyhf{{}}
\\fancyfoot{{}}
\\renewcommand{{\\headrulewidth}}{{0pt}}
\\renewcommand{{\\footrulewidth}}{{0pt}}
\\addtolength{{\\oddsidemargin}}{{-0.5in}}
\\addtolength{{\\evensidemargin}}{{-0.5in}}
\\addtolength{{\\textwidth}}{{1in}}
\\addtolength{{\\topmargin}}{{-.5in}}
\\addtolength{{\\textheight}}{{1.0in}}
\\urlstyle{{same}}
\\raggedbottom
\\raggedright
\\setlength{{\\tabcolsep}}{{0in}}
\\titleformat{{\\section}}{{\\vspace{{-4pt}}\\scshape\\raggedright\\large}}{{}}{{0em}}{{}}[\\color{{black}}\\titlerule \\vspace{{-5pt}}]
\\pdfgentounicode=1
\\newcommand{{\\resumeItem}}[1]{{\\item\\small{{{{#1 \\vspace{{-2pt}}}}}}}}
\\newcommand{{\\resumeSubheading}}[4]{{\\vspace{{-2pt}}\\item\\begin{{tabular*}}{{0.97\\textwidth}}[t]{{l@{{\\extracolsep{{\\fill}}}}r}}\\textbf{{#1}} & #2 \\\\\\textit{{\\small#3}} & \\textit{{\\small #4}} \\\\\\end{{tabular*}}\\vspace{{-7pt}}}}
\\newcommand{{\\resumeProjectHeading}}[2]{{\\item\\begin{{tabular*}}{{0.97\\textwidth}}{{l@{{\\extracolsep{{\\fill}}}}r}}\\small#1 & #2 \\\\\\end{{tabular*}}\\vspace{{-7pt}}}}
\\newcommand{{\\resumeSubHeadingListStart}}{{\\begin{{itemize}}[leftmargin=0.15in, label={{}}]}}
\\newcommand{{\\resumeSubHeadingListEnd}}{{\\end{{itemize}}}}
\\newcommand{{\\resumeItemListStart}}{{\\begin{{itemize}}}}
\\newcommand{{\\resumeItemListEnd}}{{\\end{{itemize}}\\vspace{{-5pt}}}}
\\begin{{document}}
\\begin{{center}}
    \\textbf{{\\Huge \\scshape {name}}} \\\\ \\vspace{{1pt}}
    {contact_line}
\\end{{center}}
\\section{{Education}}
  \\resumeSubHeadingListStart
    \\resumeSubheading
      {{{edu_school}}}{{{edu_location}}}
      {{{edu_degree}}}{{{edu_dates}}}
        \\resumeItemListStart
            \\resumeItem{{\\textbf{{Relevant Coursework:}} {edu_courses}}}
        \\resumeItemListEnd
  \\resumeSubHeadingListEnd
\\section{{Experience}}
  \\resumeSubHeadingListStart
{exp_latex}  \\resumeSubHeadingListEnd
\\section{{Projects}}
    \\resumeSubHeadingListStart
{proj_latex}    \\resumeSubHeadingListEnd
\\section{{Technical Skills}}
 \\begin{{itemize}}[leftmargin=0.15in, label={{}}]
    \\small{{\\item{{
{skills_latex}    }}}}
 \\end{{itemize}}
\\end{{document}}
"""


# =========================================================================
# PUBLIC API
# =========================================================================

def generate_resume(
    parsed_resume,
    jd_text: str,
    selected_experience_ids: list[str],
    selected_project_ids: list[str],
    lead_skills: list[str],
    resume_rules: str,
    similar_tech_map: dict,
    output_path: str,
    model: str = "gemini-2.5-flash",
    fallback_model: str = "gemini-2.5-flash",
    use_mock: bool = False,
) -> str | None:
    """
    Full resume generation pipeline:
    1. Tailor content with LLM (or mock)
    2. Build PDF via LaTeX (Jake's Resume format)
    Falls back to .docx if LaTeX unavailable.

    Returns path to generated file, or None on failure.
    """
    print(f"    Tailoring content...")

    if use_mock:
        tailored = tailor_resume_mock(
            parsed_resume, selected_experience_ids,
            selected_project_ids, lead_skills,
        )
    else:
        tailored = tailor_resume_with_llm(
            master_resume_text=parsed_resume.raw_text,
            jd_text=jd_text,
            selected_experience_ids=selected_experience_ids,
            selected_project_ids=selected_project_ids,
            lead_skills=lead_skills,
            parsed_resume=parsed_resume,
            resume_rules=resume_rules,
            similar_tech_map=similar_tech_map,
            model=model,
            fallback_model=fallback_model,
        )

    if not tailored:
        print(f"    ❌ Failed to tailor resume content")
        return None

    # Try LaTeX first, fall back to docx
    print(f"    Building resume (LaTeX → PDF)...")
    result = build_resume_latex(tailored, output_path)

    if result and os.path.exists(result):
        size_kb = os.path.getsize(result) / 1024
        ext = os.path.splitext(result)[1]
        print(f"    ✅ Saved: {result} ({size_kb:.0f} KB)")
        return result

    # Fallback to docx
    print(f"    LaTeX failed, falling back to .docx...")
    docx_path = os.path.splitext(output_path)[0] + ".docx"
    result = _build_resume_python_docx(tailored, docx_path)

    if result and os.path.exists(result):
        size_kb = os.path.getsize(result) / 1024
        print(f"    ✅ Saved: {result} ({size_kb:.0f} KB)")
        return result
    else:
        print(f"    ❌ Failed to build resume")
        return None
