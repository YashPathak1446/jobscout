"""
Read a PDF or DOCX resume into the structured schema the renderer expects.

The first half of Phase 2 item 10. Requiring a LaTeX master resume excluded
almost everyone; this is what lets someone arrive with the file they actually
have.

**The model fills a schema and never writes LaTeX.** `tex_renderer` owns
structure and escaping, so the worst an extraction mistake can do here is put
the right words in the wrong field — recoverable, and visible on the
confirmation screen R33 requires. Had the model been asked for markup instead,
the same mistake would produce a document that will not compile.

Text extraction is deliberately dumb. PDFs have no structure worth trusting:
two-column layouts interleave, bullets arrive as stray glyphs, and dates float
away from the roles they belong to. Rather than guess from coordinates, the
whole text goes to a model that is good at reading messy prose, and the
heuristic path below is a floor rather than a competitor.

Location: jobscout_v3/tools/resume/resume_import.py
"""

import json
import logging
import re
from pathlib import Path

logger = logging.getLogger(__name__)

SUPPORTED = {".pdf", ".docx", ".txt", ".tex"}

EXTRACTION_PROMPT = """You are reading a resume. Return ONLY a JSON object.

Extract exactly this structure, using the resume's own words. Do not invent,
summarise, reword or add anything that is not written there.

{
  "contact": {"name": "", "email": "", "phone": "", "github": "", "linkedin": ""},
  "education": [{"school": "", "degree": "", "location": "", "dates": ""}],
  "experiences": [{"company": "", "title": "", "location": "", "dates": "",
                   "bullets": ["", ""]}],
  "projects": [{"name": "", "tech": "", "dates": "", "bullets": ["", ""]}],
  "skills": {"Category Name": "comma, separated, values"}
}

Rules:
- Copy bullet text verbatim. Do not shorten or improve it.
- PDF extraction sometimes splits a word with a stray space ("W ebApp",
  "F rontend"). Repair those; do not otherwise alter wording.
- Use the LINKS FOUND section, if present, for github and linkedin. Never
  put visible link text such as "GitHub" in a URL field; leave it empty.
- "tech" is the technology list for a project, comma separated. Empty if absent.
- Keep experiences and projects in the order they appear.
- Omit any field the resume does not state. Never guess a date or a company.
- Output the JSON object alone, with no commentary and no code fence.

RESUME TEXT:
"""


# --- text extraction ---------------------------------------------------------

def extract_text(path) -> str:
    """Plain text from a resume file, whatever format it arrived in."""
    path = Path(path)
    suffix = path.suffix.lower()

    if suffix not in SUPPORTED:
        raise ValueError(
            f"Cannot read {suffix or 'a file with no extension'}; "
            f"supported: {', '.join(sorted(SUPPORTED))}")

    if suffix == ".pdf":
        return _from_pdf(path)
    if suffix == ".docx":
        return _from_docx(path)
    return path.read_text(encoding="utf-8", errors="replace")


def _from_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    text = _tidy("\n".join(pages))

    # A PDF renders link *text*, not its target, so a header showing "GitHub"
    # over a hyperlink extracts as the word "GitHub" and the URL is lost. The
    # targets live in the page annotations, so they are appended where the
    # model can see them.
    links = _pdf_links(reader)
    if links:
        text += "\n\nLINKS FOUND IN THIS DOCUMENT:\n" + "\n".join(links)
    return text


def _pdf_links(reader) -> list:
    """Every distinct http(s) target in the document's annotations."""
    found = []
    for page in reader.pages:
        for annotation in (page.get("/Annots") or []):
            try:
                action = annotation.get_object().get("/A") or {}
                uri = action.get("/URI")
            except Exception:
                continue
            uri = str(uri) if uri else ""
            if uri.startswith("http") and uri not in found:
                found.append(uri)
    return found



def _from_docx(path: Path) -> str:
    import docx

    document = docx.Document(str(path))
    lines = [p.text for p in document.paragraphs]

    # Plenty of resumes lay themselves out in tables, and skipping those loses
    # entire sections rather than a stray line.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append("  ".join(cells))

    return _tidy("\n".join(lines))


# What a T1-encoded PDF extracts as, measured from a real one.
#
# R69 found that loading \usepackage[T1]{fontenc} makes an en-dash extract as
# \x15, and rejected T1 for *this project's own* template on those grounds.
# The import path is the same finding arriving from the other side: it accepts
# PDFs from anywhere, and T1 is what most LaTeX resume advice tells people to
# load. Their file is not ours to fix, so the damage is undone on the way in.
#
# The stakes are higher than they look, because the corruption is quiet. A
# \x15 in a date is visible on the confirmation screen; "Sta Software
# Engineer" is a plausible job title, and "congurator" and "Airow" read as
# typos the owner made. Measured on one six-year resume: one corrupted job
# title, two corrupted words, and eight control characters headed for profile
# fields.
T1_ARTIFACTS = {
    "\x1b": "ff",   # Staff        -> Sta<ff>
    "\x1c": "fi",   # configurator -> con<fi>gurator
    "\x1d": "fl",   # Airflow      -> Air<fl>ow
    "\x1e": "ffi",
    "\x1f": "ffl",
    # Spaced, because PDF layout eats the space on one side: "University\x16
    # Boston" would otherwise read "University- Boston". The whitespace
    # collapse below tidies any double it creates.
    "\x15": " - ",  # en-dash, the one R69 named
    "\x16": " - ",  # em-dash
    "\x88": "",     # itemize bullet; the line break already carries it
}

# Presentation-form ligatures, which is what a PDF *without* T1 emits. Both
# encodings reach this function; only fi and fl were ever handled.
UNICODE_LIGATURES = {
    "\ufb00": "ff", "\ufb01": "fi", "\ufb02": "fl",
    "\ufb03": "ffi", "\ufb04": "ffl",
}


def _tidy(text: str) -> str:
    """Undo what PDF extraction leaves behind: ligatures, controls, space."""
    text = text.replace("\u00a0", " ")
    for artifact, plain in {**UNICODE_LIGATURES, **T1_ARTIFACTS}.items():
        text = text.replace(artifact, plain)

    # Anything left in the control range is an encoding this does not know
    # about. It must not reach a resume field either way: a job title holding
    # a raw \x0e survives every downstream check and renders as a box.
    text = re.sub(r"[\x00-\x08\x0b-\x1f\x7f-\x9f]", "", text)

    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# --- schema extraction -------------------------------------------------------

def to_schema(text: str, agent=None) -> dict:
    """
    A structured resume from raw text.

    `agent` is anything with a `_llm_json(prompt)` method — in practice the
    generation agent, so extraction rides the same backend ladder as bullet
    rewriting (R37) and inherits its caching and fallbacks. Without one, or
    when the model fails, the heuristic floor runs instead.
    """
    if agent is not None:
        try:
            parsed = agent(EXTRACTION_PROMPT + text[:24000])
            if isinstance(parsed, dict) and parsed.get("contact"):
                return _normalise(parsed)
            if parsed is None:
                # The deliberate floor: no model is configured at all. Not a
                # problem, and saying "failed" about it would be wrong.
                logger.info("No model configured; reading the resume by pattern")
            else:
                logger.warning("Extraction returned no contact block; "
                               "reading the resume by pattern instead")
        except Exception as exc:
            # A rung that should have answered did not. Distinct from having
            # no rung, and the distinction is the whole of R47: this is the
            # line that would have named the unloaded `.env` in R41 instead of
            # quietly producing a resume with no experiences.
            logger.warning(f"Resume extraction failed — {exc}. "
                           "Reading the resume by pattern instead, which will "
                           "find less. Fixing the cause is worth it.")

    return _normalise(heuristic_schema(text))


def _normalise(schema: dict) -> dict:
    """
    Fill in missing keys so callers never have to defend against them.

    `_unparsed` is carried through rather than dropped. The heuristic floor
    keeps text it could not split under that key specifically so a
    confirmation screen can show it, and building a fixed five-key dict here
    quietly threw it away — the hook existed and nothing could reach it.
    """
    out = {
        "contact": dict(schema.get("contact") or {}),
        "education": list(schema.get("education") or []),
        "experiences": list(schema.get("experiences") or []),
        "projects": list(schema.get("projects") or []),
        "skills": dict(schema.get("skills") or {}),
    }
    unparsed = schema.get("_unparsed")
    if unparsed:
        out["_unparsed"] = unparsed
    for section in ("experiences", "projects"):
        for entry in out[section]:
            entry["bullets"] = [b for b in (entry.get("bullets") or []) if b]
    return out


# --- the floor ---------------------------------------------------------------

CONTACT_PATTERNS = {
    "email": re.compile(r"[\w.+-]+@[\w-]+\.[\w.]+"),
    "phone": re.compile(r"(?:\+?\d{1,2}[\s.-]?)?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}"),
    "github": re.compile(r"https?://(?:www\.)?github\.com/[\w-]+"),
    "linkedin": re.compile(r"https?://(?:www\.)?linkedin\.com/in/[\w-]+"),
}

SECTION_HEADINGS = {
    "education": ("education",),
    "experiences": ("experience", "employment", "work history", "professional"),
    "projects": ("project", "personal project"),
    "skills": ("skill", "technical skill", "technologies"),
}


_MONTH = (r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)"
          r"[a-z]*\.?\s+")
_DATE_RANGE = re.compile(
    rf"((?:{_MONTH})?\d{{4}}\s*[-–—]\s*"
    rf"(?:(?:{_MONTH})?\d{{4}}|Present|Current))",
    re.I)
# At most three capitalised words before the comma, and no hyphen inside.
# A looser class walks backwards through "Northeastern University - Boston,
# MA" and claims the whole thing as the location, leaving the school empty —
# the greedy match is anchored on the state code, so it grows leftwards until
# it runs out of allowed characters.
_CITY_STATE = re.compile(
    r"\b([A-Z][A-Za-z.]*(?:\s+[A-Z][A-Za-z.]*){0,2},\s*[A-Z]{2})\b")
# Spelled-out qualifications and the abbreviations that cannot be mistaken for
# something else. Bare "MA" and "BA" are deliberately absent: "Boston, MA" is
# a state, and reading it as a Master of Arts put an entire university line in
# the degree field on the first resume this was tried against.
_DEGREE_WORDS = re.compile(
    r"\b(bachelor|master|associate|doctorate|doctor of|ph\.?d|"
    r"b\.?sc\.?|m\.?sc\.?|b\.s\.?|m\.s\.?|b\.a\.?|m\.a\.?|"
    r"b\.eng|m\.eng|mba|diploma)\b",
    re.I)


def _heuristic_education(lines) -> list:
    """
    One education entry with all four fields, from however many lines.

    The floor used to be `[{"school": line} for line in lines[:1]]`, which is
    two bugs in one expression. It kept the first line only — so "Bachelor of
    Science in Computer Engineering", sitting on the second line, was
    **discarded** rather than missing — and it put the whole remaining line
    into `school`, so the renderer emitted `{Northeastern University - Boston,
    MA Sep 2014 - May 2018}{}{}{}` and the degree appeared nowhere.

    The four fields are not an arity guess: `tex_renderer._education` writes
    exactly `{school}{location}{degree}{dates}` and the parser reads exactly
    those four back, so this returns the shape both ends already agree on.

    Shapes first, vocabulary second. A date range and a "City, ST" can be
    recognised by their form and are pulled out of **every** line before
    anything is classified — an earlier version classified a line as the
    degree and stopped, so a line holding all three kept all three. What
    remains is a degree if it names a qualification and part of the school
    name otherwise. Nothing is discarded: unplaced text lands in `school`,
    where it is visible on the confirmation screen (R33).
    """
    lines = [line.strip() for line in (lines or []) if line and line.strip()]
    if not lines:
        return []

    entry = {"school": "", "location": "", "degree": "", "dates": ""}
    school_parts = []

    for line in lines:
        remaining = line
        if not entry["dates"]:
            found = _DATE_RANGE.search(remaining)
            if found:
                entry["dates"] = found.group(1).strip()
                remaining = remaining.replace(found.group(1), " ")
        if not entry["location"]:
            found = _CITY_STATE.search(remaining)
            if found:
                entry["location"] = found.group(1).strip()
                remaining = remaining.replace(found.group(1), " ")

        remaining = re.sub(r"\s{2,}", " ", remaining)
        remaining = re.sub(r"^\s*[-–—|,]\s*|\s*[-–—|,]\s*$",
                           "", remaining).strip()
        if not remaining:
            continue
        if not entry["degree"] and _DEGREE_WORDS.search(remaining):
            entry["degree"] = remaining
        else:
            school_parts.append(remaining)

    entry["school"] = " ".join(school_parts).strip()
    return [entry]


def heuristic_schema(text: str) -> dict:
    """
    Contact details by pattern, sections by heading. No model.

    Honest about what this is: contact details are reliable because they have
    shapes, and everything else is a rough cut. Component boundaries, which
    bullet belongs to which role, and where a tech stack ends are exactly the
    judgements a regex cannot make. It exists so that arriving with no model
    still gets you a confirmation screen with something on it rather than an
    error, and R33's rule that every field is confirmed matters most here.
    """
    lines = [line.strip() for line in text.split("\n") if line.strip()]

    contact = {}
    for field, pattern in CONTACT_PATTERNS.items():
        match = pattern.search(text)
        if match:
            contact[field] = match.group()

    # The name is almost always the first line that is not contact details.
    for line in lines[:5]:
        if not any(p.search(line) for p in CONTACT_PATTERNS.values()) and len(line) < 60:
            contact["name"] = line
            break

    sections = _split_sections(lines)

    return {
        "contact": contact,
        "education": _heuristic_education(sections.get("education", [])),
        "experiences": [],
        "projects": [],
        "skills": ({"Skills": ", ".join(sections["skills"])}
                   if sections.get("skills") else {}),
        # Kept so a confirmation screen can show what could not be split up.
        "_unparsed": {k: v for k, v in sections.items()
                      if k in ("experiences", "projects")},
    }


def _split_sections(lines) -> dict:
    """Group lines under whichever known heading most recently appeared."""
    sections, current = {}, None

    for line in lines:
        lowered = line.lower().strip(" :")
        heading = next(
            (name for name, words in SECTION_HEADINGS.items()
             if len(lowered) < 40 and any(lowered.startswith(w) for w in words)),
            None,
        )
        if heading:
            current = heading
            sections.setdefault(current, [])
            continue
        if current:
            sections[current].append(line)

    return sections


def looks_extractable(path) -> bool:
    """Is there enough text here to bother a model with?"""
    try:
        return len(extract_text(path)) > 200
    except Exception:
        return False
